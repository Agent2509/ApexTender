from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from pydantic import BaseModel
import uuid

from api.dependencies import get_current_user_token
from database import get_db, Project

router = APIRouter()

from typing import Optional, List, Dict

class ProjectCreate(BaseModel):
    name: str
    description: Optional[str] = None

@router.get("/")
def list_projects(user: dict = Depends(get_current_user_token), db: Session = Depends(get_db)):
    return db.query(Project).filter(Project.tenant_id == user["tenant_id"]).all()

@router.post("/")
def create_project(project: ProjectCreate, user: dict = Depends(get_current_user_token), db: Session = Depends(get_db)):
    new_project = Project(
        id=str(uuid.uuid4()),
        tenant_id=user["tenant_id"],
        name=project.name,
        description=project.description
    )
    db.add(new_project)
    db.commit()
    db.refresh(new_project)
    return new_project

from fastapi import UploadFile, File
import os
import shutil
from database import DocumentMetadata
from worker import process_document_task
from supabase import create_client

@router.post("/{project_id}/documents")
async def upload_document(
    project_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user_token),
    db: Session = Depends(get_db)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDFs allowed")
        
    proj = db.query(Project).filter(Project.id == project_id, Project.tenant_id == user["tenant_id"]).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    new_doc = DocumentMetadata(
        tenant_id=user["tenant_id"],
        project_id=project_id,
        filename=file.filename,
        status="processing"
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    # os.makedirs("./uploads/", exist_ok=True)
    # temp_path = f"./uploads/{uuid.uuid4()}_{file.filename}"
    # with open(temp_path, "wb") as buffer:
    #     shutil.copyfileobj(file.file, buffer)
    
    supabase_url = os.getenv("SUPABASE_URL", "")
    supabase_key = os.getenv("SUPABASE_KEY", "")
    if not supabase_url or not supabase_key:
        raise HTTPException(status_code=500, detail="Supabase credentials not configured")
        
    supabase = create_client(supabase_url, supabase_key)
    file_bytes = await file.read()
    unique_filename = f"{uuid.uuid4()}_{file.filename}"
    
    supabase.storage.from_("documents").upload(
        path=unique_filename,
        file=file_bytes,
        file_options={"content-type": "application/pdf"}
    )
    
    public_url = supabase.storage.from_("documents").get_public_url(unique_filename)
    
    task = process_document_task.delay(public_url, user["tenant_id"], str(new_doc.id))
    
    return {"status": "success", "document_id": new_doc.id, "celery_task_id": task.id}

from qdrant_client import QdrantClient
from qdrant_client.models import Filter, FieldCondition, MatchValue
from google import genai

qdrant_url = os.getenv("QDRANT_URL")
qdrant_api_key = os.getenv("QDRANT_API_KEY")

if not qdrant_url or not qdrant_api_key:
    print("WARNING: QDRANT_URL or QDRANT_API_KEY environment variables are missing! Qdrant connection will fail.", flush=True)

qdrant_client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)

def get_query_vector(query: str):
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        client = genai.Client(api_key=api_key) if api_key else genai.Client()
        result = client.models.embed_content(
            model="models/gemini-embedding-001",
            contents=query,
            config={
                "task_type": "RETRIEVAL_QUERY",
                "output_dimensionality": 768
            }
        )
        return result.embeddings[0].values
    except Exception as e:
        print(f"Error in get_query_vector: {e}", flush=True)
        raise e

class SearchQuery(BaseModel):
    query: str
    limit: int = 5
    history: List[Dict[str, str]] = []

import os
import httpx

import json
from fastapi.responses import StreamingResponse

@router.post("/{project_id}/search")
async def search_documents(
    project_id: str,
    search_query: SearchQuery,
    user: dict = Depends(get_current_user_token),
    db: Session = Depends(get_db)
):
    try:
        proj = db.query(Project).filter(Project.id == project_id, Project.tenant_id == user["tenant_id"]).first()
        if not proj:
            raise HTTPException(status_code=404, detail="Project not found")

        query_vector = get_query_vector(search_query.query)
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"Error in search_documents setup: {e}", flush=True)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate query vector or fetch project: {str(e)}")

    try:
        # Explicitly cast to float list in case google-genai returns a different type
        float_query = [float(x) for x in query_vector]
        search_results = qdrant_client.query_points(
            collection_name="rfp_chunks",
            query=float_query,
            query_filter=Filter(
                must=[
                    FieldCondition(
                        key="project_id",
                        match=MatchValue(value=str(project_id))
                    ),
                    FieldCondition(
                        key="tenant_id",
                        match=MatchValue(value=str(user["tenant_id"]))
                    )
                ]
            ),
            limit=search_query.limit
        )
    except Exception as e:
        import traceback
        print(f"Error in Qdrant query_points: {e}", flush=True)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Qdrant Search Error: {str(e)}")

    sources = []
    context_text = ""
    for hit in search_results.points:
        chunk = hit.payload.get("text", "")[:500]
        filename = hit.payload.get("filename", "unknown")
        page_number = hit.payload.get("page_number", None)
        # Build a short snippet (first 200 chars) for citation display
        snippet = chunk[:200].strip() + ("..." if len(chunk) > 200 else "")
        sources.append({
            "chunk_text": chunk,
            "filename": filename,
            "page_number": page_number,
            "snippet": snippet,
            "document_id": hit.payload.get("document_id"),
            "similarity_score": hit.score
        })
        context_text += f"\n---\nSource: {filename} (Page {page_number})\n{chunk}"


    system_prompt = (
        "You are ApexTender, a Senior RFP Proposal Analyst. Your job is to analyze Request for Proposal (RFP) documents and provide highly accurate, concise, and structured answers to help your team win government and enterprise contracts.\n"
        "Rules:\n"
        "1. Only use the provided context. If the answer is not in the context, explicitly say: 'That information is not available in the provided documents.' Do not guess.\n"
        "2. Use professional, assertive business language.\n"
        "3. Format your answers clearly using bullet points and bold text for key terms."
    )
    user_prompt = f"Context from retrieved documents:{context_text}\n\nQuestion: {search_query.query}"

    # Build multi-turn messages
    groq_messages = [
        {"role": "system", "content": system_prompt}
    ]
    print(f"[SEARCH] Received history with {len(search_query.history)} messages", flush=True)
    for msg in search_query.history:
        if msg.get("role") in ("user", "assistant") and msg.get("content"):
            role = "user" if msg["role"] == "user" else "assistant"
            groq_messages.append({"role": role, "content": msg["content"]})
    groq_messages.append({"role": "user", "content": user_prompt})
    print(f"[SEARCH] Sending {len(groq_messages)} total messages to Groq")
    
    from groq import AsyncGroq
    
    async def generate_stream():
        try:
            api_key = os.getenv("GROQ_API_KEY", "")
            if not api_key:
                yield f"data: {json.dumps({'type': 'error', 'text': 'GROQ_API_KEY is not set'})}\n\n"
                return
                
            groq_client = AsyncGroq(api_key=api_key)
            
            response = await groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=groq_messages,
                stream=True
            )
            
            async for chunk in response:
                content = chunk.choices[0].delta.content
                if content:
                    yield f"data: {json.dumps({'type': 'chunk', 'text': content})}\n\n"
            
            # Send sources at the very end
            yield f"data: {json.dumps({'type': 'sources', 'data': sources})}\n\n"
            
        except Exception as e:
            print(f"LLM Stream Error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

    return StreamingResponse(
        generate_stream(), 
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive"
        }
    )

# --- Word Document Export ---
import io
from fastapi.responses import Response
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportRequest(BaseModel):
    content: str

@router.post("/{project_id}/export")
async def export_document(
    project_id: str,
    export_req: ExportRequest,
    user: dict = Depends(get_current_user_token),
    db: Session = Depends(get_db)
):
    proj = db.query(Project).filter(Project.id == project_id, Project.tenant_id == user["tenant_id"]).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    doc = DocxDocument()

    # Style the title
    title = doc.add_heading("RFP Generation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 64, 175)

    # Add project name as subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Project: {proj.name}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("")  # spacer

    # Add the AI-generated content
    doc.add_heading("AI-Generated Response", level=1)
    for paragraph_text in export_req.content.split("\n"):
        if paragraph_text.strip():
            p = doc.add_paragraph(paragraph_text)
            for run in p.runs:
                run.font.size = Pt(11)

    # Write to in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=Generated_Response.docx"
        }
    )

@router.delete("/{project_id}")
def delete_project(project_id: str, user: dict = Depends(get_current_user_token), db: Session = Depends(get_db)):
    proj = db.query(Project).filter(Project.id == project_id, Project.tenant_id == user["tenant_id"]).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
        
    documents = db.query(DocumentMetadata).filter(DocumentMetadata.project_id == project_id).all()
    for doc in documents:
        try:
            qdrant_client.delete(
                collection_name="rfp_chunks",
                points_selector=Filter(
                    must=[
                        FieldCondition(
                            key="filename",
                            match=MatchValue(value=doc.filename)
                        ),
                        FieldCondition(
                            key="project_id",
                            match=MatchValue(value=str(project_id))
                        )
                    ]
                )
            )
        except Exception as e:
            print(f"Error deleting vectors from Qdrant: {e}", flush=True)
            
        db.delete(doc)
        
    db.delete(proj)
    db.commit()
    return {"status": "Success", "message": "Project deleted"}
